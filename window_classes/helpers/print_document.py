from datetime import datetime
from docx import Document
from docx.shared import Inches


def print_document(filename, name, thing, color, image_filename):
    """ Создать документ о выписке потерянной вещи """
    date = str(datetime.now().date())
    document = Document()
    document.add_heading('Документ о возвращении потерянной вещи', 0)
    # Внести данные
    p = document.add_paragraph('Дата: ')
    p.add_run(date).bold = True
    p = document.add_paragraph('Имя получателя: ')
    p.add_run(name).italic = True
    p = document.add_paragraph('Предмет: ')
    p.add_run(f'"{thing}"').bold = True
    p = document.add_paragraph('Указанный цвет: ')
    p.add_run(color).bold = True
    document.add_picture(image_filename, width=Inches(5))
    # Сохранить
    document.save(filename)
