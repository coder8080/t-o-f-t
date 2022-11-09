from PyQt5.QtWidgets import QMainWindow
from .helpers.get_filename_by_id import get_filename_by_id
from .helpers.error_handlers import *
from .helpers.constants import *
from .helpers.get_filename import get_filename
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from PyQt5 import uic
import csv


class FilesWindow(QMainWindow):
    """ Окно вывода в файлы """

    def __init__(self, cur, con):
        super().__init__()
        self.cur = cur
        self.con = con
        self.initUi()

    def update_statusbar(self, message):
        self.statusbar.showMessage(message)

    def initUi(self):
        uic.loadUi('./windows/files.ui', self)
        self.update_statusbar('Ожидание выбора')
        self.csv_button.clicked.connect(self.save_csv)
        self.docx_button.clicked.connect(self.save_docx)
        self.xlsx_button.clicked.connect(self.save_xlsx)

    def fetch_items(self):
        """ Получить записи из базы данных """
        result = []
        # Получить имёна цветов
        db_colors = self.cur.execute(
            'SELECT id, name FROM colors').fetchall()
        colornames = {el[0]: el[1] for el in db_colors}
        data = self.cur.execute(
            'SELECT name, color_id, filename_id FROM things').fetchall()
        # Получить записи
        for name, color_id, filename_id in data:
            color_name = colornames[color_id]
            filename = get_filename_by_id(filename_id)
            result.append(
                {'name': name, 'color': color_name, 'filename': filename, 'filename_id': filename_id})
        return result

    def save_csv(self):
        """ Экспортировать в формат csv """
        filename = get_filename(self, 'csv', 'save')
        if not filename:
            return
        file = open(filename, mode='wt', encoding='utf-8', newline='')
        writer = csv.DictWriter(file, fieldnames=(
            'название', 'цвет', 'изображение', 'id изображения'), delimiter=';', quotechar='"', quoting=csv.QUOTE_NONNUMERIC)
        writer.writeheader()
        try:
            for item in self.fetch_items():
                writer.writerow(
                    {'название': item['name'], 'цвет': item['color'], 'изображение': item['filename'], 'id изображения': item['filename_id']})
            self.update_statusbar(FILE_GENERATED)
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            self.update_statusbar(LOADING_ERROR_MESSAGE)
        except Exception as error:
            handle_unknown_error(self, error)
            self.update_statusbar(UNKNOWN_ERROR_MESSAGE)
        finally:
            file.close()

    def save_docx(self):
        """ Экспорт в формат docx """
        filename = get_filename(self, 'docx', 'save')
        if not filename:
            return
        document = Document()
        document.add_heading('Потерянные вещи', 0)
        try:
            items = self.fetch_items()
            for item in items:
                document.add_heading(item['name'].capitalize(), 1)
                document.add_picture(item['filename'], width=Inches(3))
                document.add_paragraph(item['color'])
            if len(items) == 0:
                document.add_paragraph('Нет потерянных предметов')
            document.save(filename)
            self.update_statusbar(FILE_GENERATED)
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            self.update_statusbar(LOADING_ERROR_MESSAGE)
        except Exception as error:
            handle_unknown_error(self, error)
            self.update_statusbar(UNKNOWN_ERROR_MESSAGE)

    def save_xlsx(self):
        """ Экспорт в формат xlsx """
        filename = get_filename(self, 'xlsx', 'save')
        if not filename:
            return
        wb = Workbook()
        ws = wb.active
        ws.append(['Название', 'Цвет', 'Файл с изображением'])
        try:
            items = self.fetch_items()
            for item in items:
                arr = item.values()
                ws.append(list(arr))
            wb.save(filename)
            self.update_statusbar(FILE_GENERATED)
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            self.update_statusbar(LOADING_ERROR_MESSAGE)
        except Exception as error:
            handle_unknown_error(self, error)
            self.update_statusbar(UNKNOWN_ERROR_MESSAGE)

    def closeEvent(self, event) -> None:
        self.update_statusbar('Ожидание выбора')
