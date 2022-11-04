from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QInputDialog, QMessageBox, QTableWidgetItem, QPushButton
from PyQt5.QtGui import QPixmap
from PyQt5 import uic
from sys import argv, exit
from PIL import Image
import os
import sqlite3
import uuid0
from datetime import datetime
import csv
from docx import Document
from docx.shared import Inches

# Глобальные константы
DATABASE_FILENAME = './database.sqlite3'
LOADING_ERROR = (sqlite3.OperationalError, IndexError)
FILE_GENERATED = 'Файл сгенерирован'
LOADING_ERROR_MESSAGE = 'Ошибка при загрузке данных'
UNKNOWN_ERROR_MESSAGE = 'Ошибка'

# Вспомогательные функции


def get_filename_by_id(id: str):
    """ Возвращает имя файла по id """
    return f'./images/{id}.png'


def handle_loading_error(instance: QMainWindow, error: Exception):
    """ Выводит сообщение об ошибке загрузки """
    print(error)
    QMessageBox.critical(instance, 'Ошибка', 'Не удалось загрузить данные')


def handle_unknown_error(instance, error):
    """ Выводит сообщение о неизвестной ошибке """
    print(error)
    QMessageBox.critical(instance, 'Неизвестная ошибка',
                         'Произошла неизвестная ошибка')


def create_database():
    """ С нуля создаёт базу данных """
    file = open(DATABASE_FILENAME, mode='wb')
    file.close()
    del file
    con = sqlite3.connect('database.sqlite3')
    cur = con.cursor()
    # Создать таблицу цветов
    cur.execute(
        'CREATE TABLE colors (id INTEGER PRIMARY KEY AUTOINCREMENT, name text);').fetchall()
    # Создать таблицу вещей
    cur.execute('CREATE TABLE things (id INTEGER PRIMARY KEY AUTOINCREMENT, name text,'
                ' color_id integer, filename_id text)').fetchall()
    # Добавить базовые цвета
    cur.execute('INSERT INTO colors (name) VALUES ("красный"),'
                ' ("синий"), ("зелёный"), ("белый"), ("чёрный")').fetchall()
    # Сохраненить изменения
    con.commit()
    con.close()


def distance(a, b):
    """ Алгоритм левенштейна """
    n, m = len(a), len(b)
    if n > m:
        a, b = b, a
        n, m = m, n
    current_column = range(n+1)
    for i in range(1, m+1):
        previous_column, current_column = current_column, [i]+[0]*n
        for j in range(1, n+1):
            add, delete, change = previous_column[j] + \
                1, current_column[j-1]+1, previous_column[j-1]
            if a[j-1] != b[i-1]:
                change += 1
            current_column[j] = min(add, delete, change)

    return current_column[n]


def save_image(filename, id):
    """ Скопировать и сконвертировать изображение """
    # Создать папку если не существует
    if not os.path.exists('./images'):
        os.mkdir('./images')
    im = Image.open(filename)
    im2 = None
    # Задать новый размер изображению
    x, y = im.size
    nx, ny = 700, 450
    nratio = nx / ny
    ratio = x / y
    if ratio > nratio:
        xratio = nx / x
        height = int(xratio * y)
        im2 = im.resize((nx, height))
    elif ratio < nratio:
        yratio = ny / y
        width = int(yratio * x)
        im2 = im.resize((width, ny))
    # Скопировать изображения с изменениями
    copy_filename = get_filename_by_id(id)
    im2.save(copy_filename)


def print_document(filename, name, thing, color, image_filename):
    """ Создать документ о выписке потерянной вещи """
    date = str(datetime.now().date())
    document = Document()
    document.add_heading('Документ о возвращении потерянной вещи', 0)
    # Внести данные
    p = document.add_paragraph('Дата: ')
    p.add_run(date).bold = True
    p = document.add_paragraph('Имя получателя: ')
    p.add_run(name).italic = True
    p = document.add_paragraph('Предмет: ')
    p.add_run(f'"{thing}"').bold = True
    p = document.add_paragraph('Указанный цвет: ')
    p.add_run(color).bold = True
    document.add_picture(image_filename, width=Inches(5))
    # Сохранить
    document.save(filename)


class Logger:
    """ Инструмент сохранения записей о событиях """

    def __init__(self, filename):
        self.filename = filename
        self.file = open(filename, 'w', newline='', encoding="utf8")
        self.writer = csv.DictWriter(
            self.file, fieldnames=['время', 'имя', 'название', 'цвет'], delimiter=';', quotechar='"', quoting=csv.QUOTE_NONNUMERIC)
        self.writer.writeheader()

    def log(self, name, thing, color):
        now = str(datetime.now())
        self.writer.writerow(
            {'время': now, 'имя': name, 'название': thing, 'цвет': color})

    def close(self):
        self.file.close()


logger = Logger('TOFT.log.csv')


class FilesWindow(QMainWindow):
    """ Окно вывода в файлы """

    def __init__(self, cur, con):
        super().__init__()
        self.cur = cur
        self.con = con
        self.initUi()

    def update_statusbar(self, message):
        self.statusbar.showMessage(message)

    def initUi(self):
        uic.loadUi('./windows/files.ui', self)
        self.update_statusbar('Ожидание выбора')
        self.csv_button.clicked.connect(self.save_csv)
        self.docx_button.clicked.connect(self.save_docx)

    def get_filename(self, resolution):
        """ Получить имя файла в корректный форме """
        filename = QFileDialog.getSaveFileName(
            self, 'Получившийся файл', filter=f'{resolution}(*.{resolution})')[0]
        if not filename:
            return
        splitted = filename.split('.')
        if splitted[-1] != resolution:
            if filename[-1] != '.':
                filename += '.'
            filename += resolution
        return filename

    def fetch_items(self):
        """ Получить записи из базы данных """
        result = []
        # Получить имёна цветов
        db_colors = self.cur.execute(
            'SELECT id, name FROM colors').fetchall()
        colornames = {el[0]: el[1] for el in db_colors}
        data = self.cur.execute(
            'SELECT name, color_id, filename_id FROM things').fetchall()
        # Получить записи
        for name, color_id, filename_id in data:
            color_name = colornames[color_id]
            filename = get_filename_by_id(filename_id)
            result.append(
                {'name': name, 'color': color_name, 'filename': filename})
        return result

    def save_csv(self):
        """ Экспортировать в формат csv """
        filename = self.get_filename('csv')
        if not filename:
            return
        file = open(filename, mode='wt', encoding='utf-8', newline='')
        writer = csv.DictWriter(file, fieldnames=(
            'название', 'цвет', 'изображение'), delimiter=';', quotechar='"', quoting=csv.QUOTE_NONNUMERIC)
        writer.writeheader()
        try:
            for item in self.fetch_items():
                writer.writerow(
                    {'название': item['name'], 'цвет': item['color'], 'изображение': item['filename']})
            self.update_statusbar(FILE_GENERATED)
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            self.update_statusbar(LOADING_ERROR_MESSAGE)
        except Exception as error:
            handle_unknown_error(self, error)
            self.update_statusbar(UNKNOWN_ERROR_MESSAGE)
        finally:
            file.close()

    def save_docx(self):
        """ Экспорт в формат docx """
        filename = self.get_filename('docx')
        document = Document()
        document.add_heading('Потерянные вещи', 0)
        try:
            items = self.fetch_items()
            for item in items:
                document.add_heading(item['name'].capitalize(), 1)
                document.add_picture(item['filename'], width=Inches(3))
                document.add_paragraph(item['color'])
            if len(items) == 0:
                document.add_paragraph('Нет потерянных предметов')
            self.update_statusbar(FILE_GENERATED)
            document.save(filename)
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            self.update_statusbar(LOADING_ERROR_MESSAGE)
        except Exception as error:
            handle_unknown_error(self, error)
            self.update_statusbar(UNKNOWN_ERROR_MESSAGE)

    def closeEvent(self, event) -> None:
        self.update_statusbar('Ожидание выбора')


class ItemWindow(QMainWindow):
    """ Окно с подробной информацией о предмете """

    def __init__(self, cur, con, close_search):
        super().__init__()
        self.cur = cur
        self.con = con
        self.id = 0
        self.data = dict()
        self.filename = ''
        self.close_search = close_search
        self.initUi()

    def initUi(self):
        uic.loadUi('./windows/item.ui', self)
        self.cancel_button.clicked.connect(self.close)
        self.take_button.clicked.connect(self.take)

    def fetchData(self, id):
        """ Обновить данные """
        try:
            self.data = dict()
            result = self.cur.execute(
                f'SELECT name, filename_id, color_id FROM things WHERE id={id}').fetchall()[0]
            self.data['name'] = result[0]
            self.data['filename_id'] = result[1]
            self.data['color_id'] = result[2]
            self.data['color_name'] = self.cur.execute(
                f'SELECT name FROM colors WHERE id = {self.data["color_id"]}').fetchall()[0][0]
            self.id = id
            self.updateUi()
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            self.close()
        except Exception as error:
            handle_unknown_error(self, error)
            self.close()

    def updateUi(self):
        """ Обновить интерфейс """
        self.name_label.setText(self.data['name'])
        self.color_label.setText(self.data['color_name'])
        self.filename = get_filename_by_id(self.data['filename_id'])
        pixmap = QPixmap(self.filename)
        self.image_label.setPixmap(pixmap)

    def take(self):
        """Оформить возврат """
        name = QInputDialog.getText(
            self, 'Данные', 'Введите ваше имя')[0].strip().strip('\n')
        if not name:
            QMessageBox.critical(
                self, 'Ошибка', 'Ваши персональные данные необходимы для оформления')
            return
        try:
            # Удалить запись
            self.cur.execute(f'DELETE FROM things WHERE id = {self.id}')
            self.con.commit()
            # Создать запись в логе
            logger.log(name, self.data['name'],
                       self.data['color_name'],)
            # Напечатать документ о возврате
            print_document(
                'doc.docx', name, self.data['name'], self.data['color_name'],  self.filename)
            # Удалить ненужный файл
            os.remove(self.filename)
            # Уведомить пользователя об успехе
            QMessageBox.information(self, 'Уведомление', 'Документ сохранён в'
                                    ' файл doc.docx. Обратитесь с ним в комнату'
                                    ' забытых вещей.')
            self.close_search()
            self.close()
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
            return
        except Exception as error:
            handle_unknown_error(self, error)
            return


class SearchWindow(QMainWindow):
    """ Окно поиска по забытым вещам """

    def __init__(self, cur, con):
        super().__init__()
        self.cur = cur
        self.con = con
        self.colors = []
        self.db_colors_ids = dict()
        self.db_colors_names = dict()
        self.item_window = ItemWindow(self.cur, self.con, self.close)
        self.initUi()
        self.update_colors()

    def initUi(self):
        uic.loadUi('./windows/search.ui', self)
        self.cancel_button.clicked.connect(self.close)
        self.search_button.clicked.connect(self.search)
        self.init_table()
        self.reset_statusbar()

    def init_table(self):
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(('Название', 'Цвет', 'Просмотр'))

    def reset_name(self):
        self.name_edit.setText('')

    def reset_color_value(self):
        self.color_box.setCurrentIndex(0)

    def reset_statusbar(self):
        self.statusbar.showMessage('Ожидание действий')

    def reset_colors(self):
        self.color_box.clear()
        self.colors = []
        self.db_colors_ids = dict()

    def reset_table(self):
        self.table.setRowCount(0)

    def closeEvent(self, event):
        self.reset_name()
        self.reset_color_value()
        self.reset_statusbar()
        self.reset_table()

    def onclick_fabric(self, id):
        def result():
            self.item_window.fetchData(id)
            self.item_window.show()
        return result

    def update_colors(self):
        """ Обновить список цветов """
        self.reset_colors()
        queryColors = self.cur.execute(
            'SELECT * FROM colors').fetchall()
        for id, name in queryColors:
            self.db_colors_ids[name] = id
            self.db_colors_names[id] = name
        self.colors = ['Любой'] + [el[1] for el in queryColors]
        self.color_box.addItems(self.colors)

    def search(self):
        """ Выполнить поиск """
        name = self.name_edit.text()
        color = self.color_box.currentText()
        if color == 'Любой':
            color = None
        # Формирование запроса
        query = 'SELECT  name, color_id, id FROM things'
        if name or color:
            query += ' WHERE'
            if name:
                query += f' levenshtein(name, "{name}") < 4'
            if name and color:
                query += ' AND'
            if color:
                color_id = self.db_colors_ids[color]
                query += f' color_id = {color_id}'
        query += ';'
        result = self.cur.execute(query).fetchall()
        # Отображение данных в таблице
        self.table.setRowCount(len(result))
        for i, row in enumerate(result):
            self.table.setItem(i, 0, QTableWidgetItem(row[0]))
            self.table.setItem(i, 1, QTableWidgetItem(
                self.db_colors_names[row[1]]))
            button = QPushButton('Открыть', self)
            button.clicked.connect(self.onclick_fabric(row[2]))
            self.table.setCellWidget(i, 2, button)
        if len(result) > 0:
            self.statusbar.showMessage('Успешно')
        else:
            self.statusbar.showMessage('Записи не найдены')


class AddWindow(QMainWindow):
    """ Окно занесения нового предмета в базу """

    def __init__(self, cur, con):
        super().__init__()
        self.cur = cur
        self.con = con
        self.colors = []
        self.filename = ''
        self.initUi()
        self.update_colors()

    def initUi(self):
        uic.loadUi('./windows/add.ui', self)
        self.cancel_button.clicked.connect(self.close)
        self.add_color_button.clicked.connect(self.add_color)
        self.image_button.clicked.connect(self.open_image)
        self.submit_button.clicked.connect(self.submit)

    def reset(self):
        """ Сброс интерфейса """
        self.filename = 'Файл не выбран'
        self.update_filename()
        self.filename = None
        self.name_edit.setText('')
        self.color_box.setCurrentIndex(0)

    def closeEvent(self, e):
        self.reset()

    def add_color(self):
        """ Процесс добавления нового цвета """
        colorname = QInputDialog.getText(
            self, "Добавление цвета", "Название цвета")[0].lower()
        if not colorname:
            return
        exists = bool(self.cur.execute(
            f'SELECT * FROM colors WHERE name = "{colorname}"').fetchall())

        if exists:
            QMessageBox().warning(self, 'Добавление цвета', 'Такой цвет уже существует',
                                  QMessageBox.StandardButton.Ok)
        else:
            self.cur.execute(f'INSERT INTO colors (name) VALUES'
                             f' ("{colorname}")')
            self.con.commit()
            QMessageBox.information(self, 'Добавление цвета', 'Цвет успешно добавлен',
                                    QMessageBox.StandardButton.Ok)
            self.update_colors()
            self.color_box.setCurrentIndex(self.colors.index(colorname))

    def update_colors(self):
        """ Обновить список цветов в dropdown (combo box) """
        self.color_box.clear()
        dbColors = self.cur.execute(f'SELECT name FROM colors').fetchall()
        self.colors = [el[0] for el in dbColors]
        self.color_box.addItems(self.colors)

    def update_filename(self):
        self.image_path.setText(self.filename)

    def open_image(self):
        """ Выбрать путь до файла с изображением """
        filename = QFileDialog.getOpenFileName(
            self, 'Выберите файл с изображением', '',
            'Image Files(*.jpg , *.jpeg , *.png , *.bpm , *.webp)')[0]
        if not filename:
            QMessageBox.warning(self, 'Ошибка', 'Файл не был выбран. Скорее'
                                ' всего, вы нажали "Cancel" вместо "Open"')
            return
        self.filename = filename
        self.update_filename()

    def submit(self):
        """ Создать запись """
        name = self.name_edit.text().lower()
        filename = self.filename
        color = self.color_box.currentText()
        # Валидация полей
        if not name:
            QMessageBox.critical(
                self, 'Ошибка', 'Укажите название')
            return
        if not filename:
            QMessageBox.critical(
                self, 'Ошибка', 'Укажте имя файла с изображением')
            return
        if not color:
            QMessageBox.critical(self, 'Ошибка', 'Укажите цвет')
            return
        try:
            # Если запись с таким названием уже существует, то не создавать заново
            exists = self.cur.execute(
                f'SELECT id FROM things WHERE name = "{name}";').fetchall()
            if exists:
                QMessageBox.warning(
                    self, 'Не удалось добавить', 'Запись с таким названием уже существует',)
                return
            # Создание записи
            id = uuid0.generate()
            color_id = self.cur.execute(
                f'SELECT Id FROM colors WHERE name = "{color}"').fetchall()[0][0]
            query = f'INSERT INTO things(name, color_id, filename_id) VALUES ("{name}", {color_id}, "{id}")'
            self.cur.execute(query)
            self.con.commit()
            save_image(filename, id)
            QMessageBox.information(self, 'Уведомление', 'Запись добавлена')
            self.close()
        except LOADING_ERROR as error:
            handle_loading_error(self, error)
        except Exception as error:
            handle_unknown_error(self, error)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('./windows/main.ui', self)
        self.setupDb()
        self.initUi()

    def initUi(self):
        self.search_window = SearchWindow(self.cur, self.con)
        self.add_window = AddWindow(self.cur, self.con)
        self.files_window = FilesWindow(self.cur, self.con)

        self.search_button.clicked.connect(self.open_search_window)
        self.add_button.clicked.connect(self.open_add_window)
        self.file_button.clicked.connect(self.open_files_window)

    def setupDb(self):
        """ Установить соединение с базой данных """
        if not os.path.exists(DATABASE_FILENAME):
            create_database()
        self.con = sqlite3.connect(DATABASE_FILENAME)
        self.con.create_function('levenshtein', 2, distance)
        self.cur = self.con.cursor()

    def open_search_window(self):
        self.search_window.update_colors()
        self.search_window.show()

    def open_add_window(self):
        self.add_window.show()

    def open_files_window(self):
        self.files_window.show()

    def closeEvent(self, event):
        logger.close()


if __name__ == '__main__':
    app = QApplication(argv)
    main_window = MainWindow()
    main_window.show()
    exit(app.exec())
